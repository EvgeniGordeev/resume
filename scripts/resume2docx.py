#!/usr/bin/env python3
"""resume.json -> ATS-clean .docx, built directly with python-docx.

    python3 scripts/resume2docx.py [resume.json] [out.docx]

No pandoc and no reference document: this file owns the layout, so there is one place to change
styling. Single column, no tables, images, headers or footers -- the things ATS parsers cannot read.
Word compatibility mode is set to 15, so Word does not open the file in Compatibility Mode.

resume.json keys used here that the website ignores:
  docx              output path
  skills[].label    group label (falls back to the tier name)
  contact.phone_print / phone_print_uri   the number the exports show
  experience[].compact   render the role as a single line
  experience[].print     false = leave the role out of the exports entirely
"""
import json, os, re, sys

try:
    from docx import Document
except ImportError:
    raise SystemExit("python-docx is required: pip install python-docx")
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# --- theme -------------------------------------------------------------------
FONT      = "Calibri"
INK       = (0x1A, 0x1A, 0x1A)
GRAY      = (0x5A, 0x5A, 0x5A)
ACCENT    = (0x0F, 0x4C, 0x66)   # section headings
ACCENT_HEX = "0F4C66"
RULE_HEX  = "D9981E"             # gold rule under the contact block
LINK      = (0x1F, 0x5C, 0x99)
SEP       = "   |   "
CONTACT_SEP = "  |  "           # tighter, so the whole contact block fits on one line

MONTHS = {m: i for i, m in enumerate(
    ["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"], 1)}


def mm_yyyy(text):
    """'January 2024 - Present' -> '01/2024 - Present'; 'Oct - Dec 2007' -> '10/2007 - 12/2007'.
    A year-only range such as '2006 - 2010' is left alone."""
    t = text.strip()
    m = re.fullmatch(r'([A-Za-z]+)\.?\s*-\s*([A-Za-z]+)\.?\s+(\d{4})', t)
    if m:
        a, b, y = m.groups()
        return f"{MONTHS[a[:3].lower()]:02d}/{y} - {MONTHS[b[:3].lower()]:02d}/{y}"
    def sub(m):
        mon = m.group(1)[:3].lower()
        return f"{MONTHS[mon]:02d}/{m.group(2)}" if mon in MONTHS else m.group(0)
    return re.sub(r'([A-Za-z]+)\.?\s+(\d{4})', sub, t)


def bare(url):
    return re.sub(r'^https?://(www\.)?', '', url).rstrip('/')


def clean(s):
    return " ".join(str(s).split())


# --- docx plumbing -----------------------------------------------------------
def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin = s.bottom_margin = Inches(0.55)
    s.left_margin = s.right_margin = Inches(0.7)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.font.color.rgb = RGBColor(*INK)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.06
    set_compat_mode(doc, 15)
    return doc


def set_compat_mode(doc, val=15):
    """Word opens a file in Compatibility Mode when this is 12 (Word 2007). 15 = Word 2013+."""
    settings = doc.settings.element
    compat = settings.find(qn('w:compat'))
    if compat is None:
        compat = OxmlElement('w:compat')
        settings.append(compat)
    for cs in compat.findall(qn('w:compatSetting')):
        if cs.get(qn('w:name')) == 'compatibilityMode':
            cs.set(qn('w:val'), str(val))
            return
    cs = OxmlElement('w:compatSetting')
    cs.set(qn('w:name'), 'compatibilityMode')
    cs.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
    cs.set(qn('w:val'), str(val))
    compat.append(cs)


def _p(doc, style=None, after=4, before=0, indent=None, hang=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    if indent is not None: pf.left_indent = Inches(indent)
    if hang is not None: pf.first_line_indent = Inches(hang)
    return p


def _run(p, text, size=11, bold=False, color=None, italic=False, track=None, strike=False, caps=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.strike = strike
    r.font.all_caps = caps or None   # display-only caps: the underlying text stays readable
    if color: r.font.color.rgb = RGBColor(*color)
    if track:
        rPr = r._r.get_or_add_rPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(track)); rPr.append(sp)
    return r


def _border(p, color, sz, edge="bottom", space=4):
    pPr = p._p.get_or_add_pPr()
    bd = pPr.find(qn('w:pBdr'))
    if bd is None:
        bd = OxmlElement('w:pBdr'); pPr.append(bd)
    e = OxmlElement(f'w:{edge}')
    for k, v in (('w:val', 'single'), ('w:sz', str(sz)), ('w:space', str(space)), ('w:color', color)):
        e.set(qn(k), v)
    bd.append(e)


def hyperlink(p, url, text, size=10.5, color=LINK, bold=False):
    """A real hyperlink run: still selectable text, so parsers read the label."""
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), r_id)
    run = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rpr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '%02X%02X%02X' % color); rpr.append(col)
    if bold: rpr.append(OxmlElement('w:b'))
    run.append(rpr)
    t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve'); run.append(t)
    hl.append(run); p._p.append(hl)
    return p


def md_runs(text):
    """'~~struck~~ rest' -> [(text, strike)]. Strikethrough is the only markup carried over."""
    out, pos = [], 0
    for m in re.finditer(r'~~(.+?)~~', text):
        if m.start() > pos: out.append((text[pos:m.start()], False))
        out.append((m.group(1), True)); pos = m.end()
    if pos < len(text): out.append((text[pos:], False))
    return out


# --- building blocks ---------------------------------------------------------
def NAME(doc, name):
    p = _p(doc, after=1)
    _run(p, name, 20, bold=True, color=INK, track=24, caps=True)


def TITLE_LINE(doc, text):
    p = _p(doc, after=3)
    _run(p, text, 11.5, bold=True, color=ACCENT)


def CONTACT(doc, parts, rule=False):
    p = _p(doc, after=9 if rule else 1)
    for i, (label, url) in enumerate(parts):
        if i: _run(p, CONTACT_SEP, 9.5, color=GRAY)
        if url: hyperlink(p, url, label, size=9.5)
        else:   _run(p, label, 9.5, color=GRAY)
    if rule: _border(p, RULE_HEX, 12, space=5)


def SECTION(doc, title):
    p = _p(doc, before=13, after=5)
    _run(p, title, 12, bold=True, color=ACCENT, track=32, caps=True)
    _border(p, ACCENT_HEX, 6, space=3)


def PARA(doc, text, size=11, color=None, after=4, before=0, bold=False, italic=False):
    p = _p(doc, after=after, before=before)
    for t, strike in md_runs(text):
        _run(p, t, size, bold=bold, color=color, italic=italic, strike=strike)
    return p


def LABELED(doc, label, body, size=10.5):
    p = _p(doc, after=2)
    _run(p, label, size, bold=True, color=INK)
    _run(p, body, size)
    return p


def EMPLOYER(doc, company, url=None, before=9):
    p = _p(doc, after=0, before=before)
    if url: hyperlink(p, url, company, size=12, color=INK, bold=True)
    else:   _run(p, company, 12, bold=True, color=INK)
    return p


def ROLE(doc, title, dates):
    p = _p(doc, after=1)
    _run(p, title, 11, bold=True, color=INK)
    _run(p, SEP + dates, 10.5, color=GRAY)
    return p


def BULLET(doc, text, size=10.5):
    p = _p(doc, style="List Bullet", after=3, indent=0.19, hang=-0.13)
    _run(p, text, size)
    return p


def LINK_LINE(doc, label, links, size=10.5):
    """'Open source: a, b, c' with each item a real hyperlink."""
    p = _p(doc, after=3)
    _run(p, label, size, bold=True, color=INK)
    for i, l in enumerate(links):
        if i: _run(p, ", ", size)
        hyperlink(p, l["url"], l["label"], size=size)
    return p


# --- the document ------------------------------------------------------------
def build(data, out):
    doc = new_doc()
    c = data["contact"]

    NAME(doc, data["name"])
    tagline = " · ".join(s.strip() for s in data.get("tagline", "").split("|") if s.strip())
    TITLE_LINE(doc, data["title"] + (" · " + tagline if tagline else ""))

    phone     = c.get("phone_print") or c.get("phone")
    phone_uri = c.get("phone_print_uri") or c.get("phone_uri")
    CONTACT(doc, [(c["location"], None),
                  (phone, f"tel:{phone_uri}"),
                  (c["email"], f"mailto:{c['email']}")]
                 + [(bare(c[k]), c[k]) for k in ("linkedin", "github") if c.get(k)],
            rule=True)

    SECTION(doc, "Professional Summary")
    for para in data.get("summary", []):
        PARA(doc, clean(para))

    SECTION(doc, "Technical Skills")
    for tier in data.get("skills", []):
        LABELED(doc, (tier.get("label") or tier.get("tier", "Skills")) + ":  ", ", ".join(tier["items"]))

    SECTION(doc, "Professional Experience")
    for job in [j for j in data.get("experience", []) if j.get("print") is not False]:
        if job.get("compact"):
            p = _p(doc, after=3, before=4)
            _run(p, job["role"] + ", ", 10.5, bold=True, color=INK)
            if job.get("url"): hyperlink(p, job["url"], job["company"], size=10.5, color=GRAY)
            else:              _run(p, job["company"], 10.5, color=GRAY)
            _run(p, SEP + mm_yyyy(job["date"]), 10.5, color=GRAY)
            duties = "; ".join(clean(d).rstrip('.') for d in job.get("duties", []))
            if duties: _run(p, ". " + duties, 10.5)
            continue
        EMPLOYER(doc, job["company"], job.get("url"))
        ROLE(doc, job["role"], mm_yyyy(job["date"]))
        if job.get("duties"):
            # job["heading"] ("In charge of", "Responsibilities") is a website label, not printed here
            PARA(doc, "; ".join(clean(d).rstrip('.') for d in job["duties"]) + ".",
                 size=10.5, after=7, before=2)
        for proj in job.get("projects") or []:
            links = proj.get("highlights_as_links") or []
            if links and not proj.get("highlights") and not proj.get("description"):
                LINK_LINE(doc, proj["name"] + ":  ", links)
                continue
            PARA(doc, proj["name"], size=10.5, bold=True, after=1, before=3)
            if proj.get("description"):
                PARA(doc, clean(proj["description"]), size=10.5, color=GRAY, after=3)
            for h in proj.get("highlights") or []:
                BULLET(doc, clean(h))
            for l in links:
                p = _p(doc, style="List Bullet", after=3, indent=0.19, hang=-0.13)
                hyperlink(p, l["url"], l["label"])
            env = proj.get("env") or {}
            bits = [f"Tools: {env['tools']}" if env.get("tools") else None,
                    f"Tech: {env['techs']}"  if env.get("techs") else None]
            bits = [b for b in bits if b]
            if bits:
                PARA(doc, SEP.join(bits), size=9.5, color=GRAY, after=5)

    SECTION(doc, "Education")
    for ed in data.get("education", []):
        p = _p(doc, after=3)
        _run(p, ed["school"], 11, bold=True, color=INK)
        _run(p, SEP + mm_yyyy(ed["date"]), 10.5, color=GRAY)
        if ed.get("major"):
            _run(p, SEP + clean(ed["major"]), 10.5, color=GRAY)

    if data.get("certifications"):
        SECTION(doc, "Certifications")
        for cert in data["certifications"]:
            p = _p(doc, after=2)
            # certification names are not bold: the section heading already frames them
            if cert.get("url"): hyperlink(p, cert["url"], cert["name"], size=10.5, color=INK)
            else:               _run(p, cert["name"], 10.5, color=INK)
            _run(p, SEP + cert["issuer"] + SEP + mm_yyyy(cert["date"]), 10.5, color=GRAY)

    if data.get("languages"):
        SECTION(doc, "Languages")
        # one line, no proficiency percentages -- the bars belong on the website
        PARA(doc, SEP.join(l["name"] for l in data["languages"]), size=10.5)

    doc.core_properties.author = data["name"]
    doc.core_properties.last_modified_by = data["name"]
    for k in ("title", "subject", "keywords", "comments", "category"):
        setattr(doc.core_properties, k, "")
    doc.save(out)
    return out


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, "resume.json")
    data = json.load(open(src))
    out = sys.argv[2] if len(sys.argv) > 2 else os.path.join(
        ROOT, data.get("docx") or f"downloads/{data['name'].replace(' ', '_')}_CV.docx")
    if not os.path.isabs(out): out = os.path.join(ROOT, out)
    os.makedirs(os.path.dirname(out), exist_ok=True)
    print("Generated " + os.path.relpath(build(data, out), ROOT))
